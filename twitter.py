from selenium import webdriver
from selenium.webdriver.common.keys import Keys
from bs4 import BeautifulSoup as soup
import time
from docx import Document

####### LINK #########

# pip install selenium
# pip install python-docx
# pip install beautifulsoup4

# Selenium Documentation
# https://selenium-python.readthedocs.io/

# Chrome Driver
# https://sites.google.com/a/chromium.org/chromedriver/downloads

# BeautifulSoup Documentation
# https://www.crummy.com/software/BeautifulSoup/bs4/doc/

# Docx Documentation
# https://python-docx.readthedocs.io/en/latest/


opt = webdriver.ChromeOptions()
opt.add_argument('headless') # Don't open the browser

driver = webdriver.Chrome('C:\\Python37\\chromedriver.exe',options=opt)

pixel = 0

def HashTag(keyword,news=40):

	global pixel

	url = 'https://twitter.com/hashtag/' + keyword

	driver.get(url)

	time.sleep(3)
	totalpage =  news // 20 # per scrolling is 20 news
	for i in range(totalpage):
		driver.execute_script("window.scrollTo(0, {})".format(pixel))
		time.sleep(3)
		pixel = pixel + 10000
		#pixel += 10000

	page_html = driver.page_source

	data = soup(page_html, 'html.parser')

	tweetext = data.findAll('p',{'class':'TweetTextSize'})
	tweetuser = data.findAll('a',{'class':'account-group js-account-group js-action-profile js-user-profile-link js-nav'})

	#print(tweetuser)
	result = []

	for i,tw in enumerate(zip(tweetuser,tweetext)):
		print(i+1)
		user = tw[0].text
		user = user.replace('\n','')
		user = user.replace('\u200f\xa0',' ')
		print(user)
		print('https://twitter.com' + tw[0]['href'])
		url = 'https://twitter.com' + tw[0]['href']
		print(tw[1].text)
		print('----------')
		result.append({'user':user,'url':url,'tweet':tw[1].text})

	driver.close()

	return result

## Create Docx Report
document = Document()


########### SINGLE HASHTAG ###########
news1 = HashTag('หนุ่มแว่นหัวร้อน',40)

document.add_heading('หนุ่มแว่นหัวร้อน', 0)
p = document.add_paragraph(' ')

for nw in news1:
	p.add_run(nw['user']).bold = True
	p.add_run('\n')
	p.add_run(nw['url'] + '\n')
	p.add_run(nw['tweet'] + '\n\n\n')
########### SINGLE HASHTAG ###########

document.save('Twitter HashTag.docx')
print('Done!')



